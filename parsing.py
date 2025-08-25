import io
import os
import re
import tempfile
from pathlib import Path
from typing import Any, Dict

import docx
import fitz
from fastapi import UploadFile, HTTPException

from config import logger


def extract_text_from_pdf_bytes(pdf_bytes: bytes) -> str:
    full_text = []
    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception as e:
        logger.error(f"Error opening PDF: {e}")
        raise HTTPException(status_code=400, detail="Invalid PDF file.")

    for page in doc:
        page_texts = []

        simple_text = page.get_text().strip()
        if simple_text:
            page_texts.append(simple_text)

        blocks = page.get_text("blocks")
        if blocks:
            sorted_blocks = sorted(blocks, key=lambda b: (round(b[1] / 10) * 10, b[0]))
            block_text = "\n".join([b[4].strip() for b in sorted_blocks if len(b) > 4 and b[4].strip()])
            if block_text:
                page_texts.append(block_text)

        if page_texts:
            best_text = max(page_texts, key=len)
            full_text.append(best_text)

    combined_text = "\n\n".join(full_text)
    combined_text = re.sub(r"\r\n?", "\n", combined_text)
    combined_text = re.sub(r"\n{3,}", "\n\n", combined_text)
    combined_text = re.sub(r"[ \t]+", " ", combined_text)
    return combined_text


def extract_text_from_docx_bytes(docx_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(docx_bytes)
        tmp.flush()
        tmp_path = tmp.name

    try:
        document = docx.Document(tmp_path)

        paragraphs = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        tables = []
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    tables.append(" | ".join(row_text))

        all_text = paragraphs + tables
    finally:
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    return "\n".join(all_text)


def parse_resume_file(upload_file: UploadFile) -> Dict[str, Any]:
    filename = upload_file.filename or "uploaded_resume"
    fname_low = filename.lower()

    if not (fname_low.endswith(".pdf") or fname_low.endswith(".docx")):
        raise HTTPException(
            status_code=400,
            detail=f"Invalid file type for '{filename}'. Only PDF and DOCX files are allowed."
        )

    contents = upload_file.file.read()
    if len(contents) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")

    try:
        if fname_low.endswith(".pdf"):
            text = extract_text_from_pdf_bytes(contents)
        else:
            text = extract_text_from_docx_bytes(contents)
    except Exception as e:
        logger.error(f"Error processing file {filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

    if not text or len(text.strip()) < 50:
        raise HTTPException(status_code=400, detail="Could not extract meaningful text from the resume.")

    return {"text": text, "filename": filename}
